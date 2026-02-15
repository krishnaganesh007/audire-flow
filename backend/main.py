import asyncio
import glob
import os
import io
import uuid
import re
import mammoth
import fitz # PyMuPDF
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import List, Dict, Any
from docx import Document
from fastapi.staticfiles import StaticFiles

app = FastAPI()

# CORS configuration
origins = [
    "http://localhost:5173",
    "http://localhost:5174",
    "http://127.0.0.1:5173",
    "http://127.0.0.1:5174",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

PROMPTS_DIR = os.path.join(os.path.dirname(__file__), "prompts")
TEMP_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "temp"))
EXPORT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "exports"))
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

app.mount("/exports", StaticFiles(directory=EXPORT_DIR), name="exports")

def load_prompts():
    prompts = {}
    prompt_files = glob.glob(os.path.join(PROMPTS_DIR, "*.md"))
    prompt_files.sort()
    for fpath in prompt_files:
        basename = os.path.basename(fpath)
        with open(fpath, "r") as f:
            prompts[basename] = f.read()
    return prompts

loaded_prompts = load_prompts()

class Finding(BaseModel):
    id: str
    original: str
    refined: str
    stage: int
    status: str
    decision: str # 'approved' | 'rejected' | 'pending'

class DocumentResponse(BaseModel):
    document_html: str
    findings_map: Dict[str, Finding]

class ExportRequest(BaseModel):
    original_filename: str
    export_filename: str = ""
    export_format: str = "docx"  # "docx" or "pdf"
    approved_findings: Dict[str, str]

@app.get("/")
async def root():
    return {"message": "Audire Flow Backend is running"}

async def refine_finding(finding_obj):
    # Simulate LLM latency
    await asyncio.sleep(2.0) 
    
    refined_text = finding_obj["original"] + " [Refined by AI]"
    finding_obj["refined"] = refined_text
    finding_obj["stage"] = 6
    finding_obj["status"] = "done"
    return finding_obj

@app.post("/process-document")
async def process_document(file: UploadFile = File(...)):
    if not (file.filename.endswith('.docx') or file.filename.endswith('.pdf')):
        raise HTTPException(status_code=400, detail="Only .docx and .pdf files are supported.")

    content = await file.read()
    
    temp_id = str(uuid.uuid4())
    temp_path = os.path.join(TEMP_DIR, f"{temp_id}_{file.filename}")
    with open(temp_path, "wb") as f:
        f.write(content)

    findings_map = {}
    
    if file.filename.endswith('.pdf'):
        import re
        doc_pdf = fitz.open(temp_path)
        html_content = ""
        findings_map = {}
        last_finding_col_idx = -1
        
        def extract_formatted(page, rect):
            """Helper to extract formatted HTML snippet from a specific rectangle."""
            # Expand rect slightly to avoid cutting off text
            expanded_rect = fitz.Rect(rect) + (-1, -1, 1, 1)
            p_dict = page.get_text("dict", clip=expanded_rect)
            snippet = ""
            for b_dict in p_dict.get("blocks", []):
                if b_dict["type"] != 0: continue
                # Blocks often have multiple lines, each with spans
                for l in b_dict.get("lines", []):
                    for s in l.get("spans", []):
                        txt = s["text"]
                        if not txt.strip(): continue
                        f = s["flags"]
                        styles = []
                        if f & 2**4: styles.append("font-weight:bold") # Bold
                        if f & 2**1: styles.append("font-style:italic") # Italic
                        style_attr = f' style="{";".join(styles)}"' if styles else ""
                        snippet += f'<span{style_attr}>{txt}</span>'
                    snippet += " "
            return snippet.strip()

        for page_num, page in enumerate(doc_pdf):
            # 1. Detect Tables on this page
            tab_finder = page.find_tables(horizontal_strategy="lines", vertical_strategy="lines")
            tabs = list(tab_finder)
            tab_rects = [fitz.Rect(t.bbox) for t in tabs]
            
            # 2. Extract Generic Layout
            page_dict = page.get_text("dict")
            blocks = page_dict["blocks"]
            blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0])) # Sort by Y, then X
            
            processed_tabs_on_page = set()
            
            for b_idx, b in enumerate(blocks):
                if b["type"] != 0: continue # Skip images
                bbox = fitz.Rect(b["bbox"])
                
                # Check if this block is inside any table on this page
                current_table_idx = -1
                for i, rect in enumerate(tab_rects):
                    # Only count as in-table if the majority of the block is inside
                    intersection = bbox & rect
                    if intersection.get_area() > 0.5 * bbox.get_area():
                        current_table_idx = i
                        break
                
                if current_table_idx != -1:
                    # BLOCK IS INSIDE A TABLE
                    if current_table_idx not in processed_tabs_on_page:
                        # Render the whole table once
                        tab = tabs[current_table_idx]
                        table_data = tab.extract()
                        if not table_data:
                            processed_tabs_on_page.add(current_table_idx)
                            continue
                        
                        html_content += "<table border='1' style='border-collapse: collapse; margin: 20px 0; width: 100%; font-family: inherit;'>"
                        
                        # Column & Header Detection
                        finding_col_idx = -1
                        # Check if first row is a header
                        first_row_text = " ".join([str(c or "").lower() for c in table_data[0]])
                        
                        # Use same keywords as DOCX path for consistency
                        header_keywords = ["finding", "issue", "observation", "description"]
                        is_header_row = any(k in first_row_text for k in header_keywords)
                        
                        # Fallback: Check second row if first row is empty or doesn't look like a header
                        if not is_header_row and len(table_data) > 1:
                            second_row_text = " ".join([str(c or "").lower() for c in table_data[1]])
                            if any(k in second_row_text for k in header_keywords):
                                is_header_row = True # We use the first row as headers for simplicity if hits
                        
                        if is_header_row:
                            for c_idx, cell_val in enumerate(table_data[0]):
                                norm = (cell_val or "").lower()
                                # Find the FIRST column with finding keywords, then break
                                if any(k in norm for k in header_keywords):
                                    finding_col_idx = c_idx
                                    print(f"[DEBUG] Selected finding column {c_idx}: '{cell_val}'")
                                    break  # Use the first matching column, don't keep overwriting
                            last_finding_col_idx = finding_col_idx
                            print(f"[DEBUG] Final finding_col_idx for table: {finding_col_idx}")
                        else:
                            # Use inherited column if continuation
                            finding_col_idx = last_finding_col_idx
                            
                        # Default fallback
                        if finding_col_idx == -1 and len(table_data[0]) > 1:
                            finding_col_idx = 1
                            
                        # Render Rows
                        num_cols = len(table_data[0]) if table_data else 0
                        for r_idx, row in enumerate(table_data):
                            html_content += "<tr>"
                            for c_idx, cell_text_raw in enumerate(row):
                                # Get precise cell HTML with formatting
                                cell_obj = tab.cells[r_idx * num_cols + c_idx]
                                cell_html = ""
                                if cell_obj:
                                    cell_html = extract_formatted(page, fitz.Rect(cell_obj))
                                
                                # Fallback if empty or failed
                                if not cell_html.strip():
                                    cell_html = (cell_text_raw or "").strip()
                                
                                # Styling
                                tag = "th" if (r_idx == 0 and is_header_row) else "td"
                                bg = "background-color: #f3f3f3;" if (r_idx == 0 and is_header_row) else ""
                                style = f"padding: 10px; border: 1px solid #c0c0c0; vertical-align: top; {bg}"
                                
                                # Create Finding if in finding column
                                # Use cell_text_raw (plain text from PyMuPDF) like DOCX uses cell.text
                                plain_text_for_matching = (cell_text_raw or "").strip()
                                plain_text_for_matching = " ".join(plain_text_for_matching.split())

                                if (not (r_idx == 0 and is_header_row)) and c_idx == finding_col_idx and len(plain_text_for_matching) > 10:
                                    fid = f"pdf_p{page_num}_t{current_table_idx}_r{r_idx}_{uuid.uuid4().hex[:4]}"
                                    print(f"[DEBUG] Extracted finding from row {r_idx}, col {c_idx}: '{plain_text_for_matching[:100]}...'")
                                    findings_map[fid] = {
                                        "id": fid,
                                        "original": plain_text_for_matching,
                                        "refined": plain_text_for_matching,
                                        "stage": 0, "status": "processing", "decision": "pending"
                                    }
                                
                                html_content += f"<{tag} style='{style}'>{cell_html}</{tag}>"
                            html_content += "</tr>"
                        
                        html_content += "</table>"
                        processed_tabs_on_page.add(current_table_idx)
                else:
                    # BLOCK IS REGULAR TEXT
                    block_html = ""
                    for line in b.get("lines", []):
                        line_content = ""
                        for s in line.get("spans", []):
                            txt = s["text"]
                            if not txt.strip(): continue
                            f = s["flags"]
                            style = []
                            if f & 2**4: style.append("font-weight:bold")
                            if f & 2**1: style.append("font-style:italic")
                            
                            sz = s["size"]
                            if sz > 15: tag = "h1"
                            elif sz > 13: tag = "h2"
                            elif sz > 11: tag = "h3"
                            else: tag = "span"
                            
                            s_attr = f' style="{";".join(style)}"' if style else ""
                            if tag.startswith("h"):
                                line_content += f"<{tag}{s_attr}>{txt}</{tag}>"
                            else:
                                line_content += f"<span{s_attr}>{txt}</span>"
                        
                        if line_content:
                            # Fallback finding extraction outside tables
                            clean_txt = re.sub('<[^<]+?>', '', line_content).strip()
                            clean_txt = " ".join(clean_txt.split())
                            if any(clean_txt.lower().startswith(k) for k in ["finding:", "issue:", "gap:"]):
                                fid = f"pdf_txt_p{page_num}_b{b_idx}_{uuid.uuid4().hex[:4]}"
                                findings_map[fid] = {
                                    "id": fid, "original": clean_txt, "refined": clean_txt,
                                    "stage": 0, "status": "processing", "decision": "pending"
                                }
                            block_html += f"<div>{line_content}</div>"
                    
                    if block_html:
                        html_content += f"<div style='margin-bottom: 12px;'>{block_html}</div>"

        document_html = html_content
        
        # Refine findings in parallel
        tasks = [refine_finding(f) for f in findings_map.values()]
        if tasks:
            await asyncio.gather(*tasks)

        return {
            "document_html": document_html,
            "findings_map": findings_map
        }

    # DOCX path
    with open(temp_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        document_html = result.value
        
    # Extract Findings using python-docx
    doc = Document(temp_path)
    findings_map = {}
    tasks = []
    
    for table_idx, table in enumerate(doc.tables):
        header_row = table.rows[0].cells
        finding_col_idx = -1
        
        for i, cell in enumerate(header_row):
            text = cell.text.lower()
            if any(key in text for key in ["finding", "issue", "observation", "description"]):
                finding_col_idx = i
                break
        
        if finding_col_idx != -1:
            for row_idx, row in enumerate(table.rows[1:], start=1):
                cell = row.cells[finding_col_idx]
                original_text = cell.text.strip()
                if original_text:
                    fid = f"table_{table_idx}_row_{row_idx}_col_{finding_col_idx}"
                    finding_obj = {
                        "id": fid,
                        "original": original_text,
                        "refined": original_text, 
                        "stage": 0,
                        "status": "processing",
                        "decision": "pending"
                    }
                    findings_map[fid] = finding_obj
                    tasks.append(refine_finding(finding_obj))

    if tasks:
        await asyncio.gather(*tasks)
    
    return {
        "document_html": document_html,
        "findings_map": findings_map
    }

@app.post("/export-document")
async def export_document(request: ExportRequest):
    temp_files = glob.glob(os.path.join(TEMP_DIR, f"*_{request.original_filename}"))
    if not temp_files:
        raise HTTPException(status_code=404, detail="Original document not found. Please re-upload.")

    latest_temp = max(temp_files, key=os.path.getmtime)
    doc = Document(latest_temp)

    for table_idx, table in enumerate(doc.tables):
        header_row = table.rows[0].cells
        finding_col_idx = -1
        for i, cell in enumerate(header_row):
            text = cell.text.lower()
            if any(key in text for key in ["finding", "issue", "observation", "description"]):
                finding_col_idx = i
                break
        
        if finding_col_idx != -1:
            for row_idx, row in enumerate(table.rows[1:], start=1):
                fid = f"table_{table_idx}_row_{row_idx}_col_{finding_col_idx}"
                if fid in request.approved_findings:
                    table.rows[row_idx].cells[finding_col_idx].text = request.approved_findings[fid]

    export_base = request.export_filename if request.export_filename else f"refined_{request.original_filename.replace('.docx','').replace('.pdf','')}"
    docx_filename = f"{export_base}.docx"
    docx_path = os.path.join(EXPORT_DIR, docx_filename)
    doc.save(docx_path)

    if request.export_format == "pdf":
        import subprocess
        soffice_path = "/opt/homebrew/bin/soffice"
        if not os.path.exists(soffice_path): soffice_path = "soffice"
            
        try:
            subprocess.run([
                soffice_path, "--headless", "--convert-to", "pdf",
                "--outdir", EXPORT_DIR, docx_path
            ], check=True, timeout=30)
            export_filename = f"{export_base}.pdf"
        except:
            export_filename = docx_filename
    else:
        export_filename = docx_filename

    return {
        "download_url": f"http://localhost:8000/exports/{export_filename}",
        "filename": export_filename
    }

@app.get("/prompts")
async def get_prompts():
    return loaded_prompts
